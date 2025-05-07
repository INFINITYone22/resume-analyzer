import PyPDF2
from docx import Document
from typing import BinaryIO, Dict, Any, List
import re
import io

def extract_text_from_pdf(file: BinaryIO) -> str:
    """Extract text from a PDF file."""
    # Need to seek to beginning in case file was already read
    file.seek(0)
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text += page_text + "\n"
    return text.strip()

def extract_text_from_docx(file: BinaryIO) -> str:
    """Extract text from a DOCX file."""
    file.seek(0)
    doc = Document(file)
    
    # Extract text from paragraphs
    paragraphs_text = [para.text for para in doc.paragraphs]
    
    # Extract text from tables
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            tables_text.append(" | ".join(row_text))
    
    # Combine all text
    all_text = "\n".join(paragraphs_text + tables_text)
    return all_text.strip()

def extract_resume_text(file: BinaryIO, filename: str) -> str:
    """Extract text from a resume file (PDF or DOCX)."""
    # Make a copy of the file in memory to avoid consuming it
    file_copy = io.BytesIO(file.getvalue())
    
    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_copy)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(file_copy)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

def extract_resume_sections(text: str) -> Dict[str, str]:
    """
    Attempt to extract common resume sections like education, experience, skills, etc.
    This is a simple heuristic-based approach and won't work for all resumes.
    """
    # Common section headers in resumes
    section_headers = [
        "education", "experience", "work experience", "employment", 
        "skills", "technical skills", "professional skills",
        "projects", "certifications", "achievements", "awards",
        "publications", "languages", "interests", "hobbies",
        "volunteer", "references", "summary", "objective", "profile"
    ]
    
    # Dictionary to store sections
    sections = {}
    
    # Split text into lines and process
    lines = text.split('\n')
    current_section = "header"  # Default section for content before any section header
    section_content = []
    
    for line in lines:
        line = line.strip()
        if not line:  # Skip empty lines
            continue
            
        # Check if this line is a section header
        is_header = False
        header_match = None
        
        for header in section_headers:
            # Case-insensitive match for section headers
            # Match full words only (e.g., "Skills:" but not "Technical Skills")
            pattern = r'^\s*(' + re.escape(header) + r')[\s:]*$'
            match = re.search(pattern, line.lower())
            if match:
                is_header = True
                header_match = match.group(1)
                break
        
        if is_header and header_match:
            # Save the previous section
            if section_content:
                sections[current_section] = '\n'.join(section_content)
                section_content = []
            current_section = header_match
        else:
            # Add line to current section
            section_content.append(line)
    
    # Save the last section
    if section_content:
        sections[current_section] = '\n'.join(section_content)
    
    return sections 